from docx import Document
from io import BytesIO


def create_docx(report_text, project_title):
    """
    Converts generated report text into a DOCX file.
    """

    document = Document()

    document.add_heading(
        project_title if project_title else "Generated Project Report",
        0
    )

    for line in report_text.split("\n"):
        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            document.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            document.add_heading(line.replace("### ", ""), level=3)

        elif line.startswith("- "):
            document.add_paragraph(line.replace("- ", ""), style="List Bullet")

        elif line.startswith("* "):
            document.add_paragraph(line.replace("* ", ""), style="List Bullet")

        else:
            clean_line = (
                line.replace("**", "")
                .replace("__", "")
                .replace("`", "")
            )

            document.add_paragraph(clean_line)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream